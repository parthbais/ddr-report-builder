from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os


# ── Helpers ──────────────────────────────────────────────────────────────────

def _set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _add_horizontal_rule(doc, color="2E7D32"):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def _heading(doc, text: str, bold=True, color="1A237E", size=14):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    return p


def _body(doc, text: str, size=11):
    p = doc.add_paragraph(str(text))
    for run in p.runs:
        run.font.size = Pt(size)
    p.paragraph_format.space_after = Pt(6)
    return p


def _label_value(doc, label: str, value: str, size=10.5):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    r1 = p.add_run(f"{label}: ")
    r1.bold = True
    r1.font.size = Pt(size)
    r1.font.color.rgb = RGBColor.from_string("1A237E")
    r2 = p.add_run(str(value) if value else "Not Available")
    r2.font.size = Pt(size)
    return p


def _bullet(doc, text: str, size=11):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(str(text))
    run.font.size = Pt(size)
    return p


def _find_images_for_area(images, area_name, max_count=2, min_bytes=5000):
    """Match images to an area by nearest_heading keyword overlap. Skip tiny icons."""
    if not images:
        return []

    stop = {'the', 'a', 'an', 'of', 'in', 'for', 'and', 'or', 'is', 'are', 'at',
            'to', 'by', 'as', 'on', 'its', 'with', 'from', 'not', 'no', 'this'}
    area_words = {w for w in area_name.lower().split() if w not in stop and len(w) > 2}

    scored = []
    for img in images:
        path = img.get("path", "")
        if not os.path.exists(path):
            continue
        if os.path.getsize(path) < min_bytes:
            continue
        heading = img.get("nearest_heading", "").lower()
        score = sum(1 for w in area_words if w in heading)
        if score > 0:
            scored.append((score, img))

    scored.sort(key=lambda x: x[0], reverse=True)
    return [img for _, img in scored[:max_count]]


def _insert_image(doc, img_path: str, caption: str = "", width_inches: float = 4.5):
    """Insert an image centered with an optional italic caption below."""
    try:
        p = doc.add_paragraph()
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = p.add_run()
        run.add_picture(img_path, width=Inches(width_inches))

        if caption:
            cap = doc.add_paragraph()
            cap.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            cr = cap.add_run(caption)
            cr.italic = True
            cr.font.size = Pt(9)
            cr.font.color.rgb = RGBColor.from_string("555555")
    except Exception as e:
        p = doc.add_paragraph(f"[Image not available: {e}]")
        p.runs[0].italic = True


# ── Main builder ─────────────────────────────────────────────────────────────

def create_ddr_docx(
    ddr_json: dict,
    inspection_images: list = None,
    thermal_images: list = None,
    assets_dir: str = "assets",
    output_path: str = "output/Final_DDR.docx",
) -> str:
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    inspection_images = inspection_images or []
    thermal_images    = thermal_images or []

    doc = Document()

    # ── Margins ──────────────────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin    = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2)

    # ══════════════════════════════════════════════════════════════════════════
    # COVER PAGE
    # ══════════════════════════════════════════════════════════════════════════
    doc.add_paragraph()
    brand = doc.add_paragraph()
    brand.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    r = brand.add_run("UrbanRoof Private Limited")
    r.bold = True; r.font.size = Pt(22)
    r.font.color.rgb = RGBColor.from_string("2E7D32")

    doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    r2 = title.add_run("Detailed Diagnostic Report (DDR)")
    r2.bold = True; r2.font.size = Pt(26)
    r2.font.color.rgb = RGBColor.from_string("1A237E")

    sub = doc.add_paragraph()
    sub.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    r3 = sub.add_run("Comprehensive Building Health Assessment")
    r3.font.size = Pt(13); r3.italic = True
    r3.font.color.rgb = RGBColor.from_string("555555")

    _add_horizontal_rule(doc, "2E7D32")
    doc.add_paragraph()

    # Property info table on cover
    pi = ddr_json.get("property_info", {})
    cover_fields = [
        ("Client Name",      pi.get("client_name", "N/A")),
        ("Address",          pi.get("address", "N/A")),
        ("Inspection Date",  pi.get("inspection_date", "N/A")),
        ("Inspected By",     pi.get("inspected_by", "N/A")),
        ("Property Type",    pi.get("property_type", "N/A")),
        ("No. of Floors",    pi.get("floors", "N/A")),
        ("Age of Property",  f"{pi.get('age_years','N/A')} years"),
        ("Previous Repairs", pi.get("previous_repairs", "N/A")),
        ("Tools Used",       pi.get("tools_used", "N/A")),
    ]
    cover_table = doc.add_table(rows=len(cover_fields), cols=2)
    cover_table.style = "Table Grid"
    for i, (k, v) in enumerate(cover_fields):
        row = cover_table.rows[i]
        _set_cell_bg(row.cells[0], "E8F5E9")
        row.cells[0].text = k
        row.cells[0].paragraphs[0].runs[0].bold = True
        row.cells[0].paragraphs[0].runs[0].font.size = Pt(10)
        row.cells[1].text = str(v)
        row.cells[1].paragraphs[0].runs[0].font.size = Pt(10)

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # TABLE OF CONTENTS (manual)
    # ══════════════════════════════════════════════════════════════════════════
    toc_title = doc.add_paragraph()
    toc_title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    r = toc_title.add_run("Table of Content")
    r.bold = True; r.font.size = Pt(16); r.font.underline = True
    r.font.color.rgb = RGBColor.from_string("1A237E")
    _add_horizontal_rule(doc, "2E7D32")

    obs_list = ddr_json.get("area_wise_observations", [])
    toc_items = [
        ("SECTION 1", "INTRODUCTION", "3"),
        ("  1.1",     "Background", "3"),
        ("  1.2",     "Objective of the Health Assessment", "3"),
        ("  1.3",     "Scope of Work", "3"),
        ("SECTION 2", "GENERAL INFORMATION", "4"),
        ("  2.1",     "Client & Inspection Details", "4"),
        ("  2.2",     "Description of Site", "4"),
        ("SECTION 3", "VISUAL OBSERVATION AND READINGS", "5"),
        ("  3.1",     "Executive Summary & Issue Overview", "5"),
        ("  3.2",     "Sources of Leakage / Root Cause", "5"),
    ]
    for idx, obs in enumerate(obs_list, start=3):
        toc_items.append((f"  3.{idx}", obs.get("area", f"Area {idx}"), str(5 + idx)))
    toc_items += [
        ("SECTION 4", "SEVERITY ASSESSMENT", str(6 + len(obs_list))),
        ("SECTION 5", "RECOMMENDED ACTIONS", str(7 + len(obs_list))),
        ("SECTION 6", "WATERPROOFING & STRUCTURAL RECOMMENDATIONS", str(8 + len(obs_list))),
        ("SECTION 7", "ADDITIONAL NOTES", str(9 + len(obs_list))),
    ]

    for num, title_text, page in toc_items:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        is_section = not num.startswith("  ")
        r_num = p.add_run(f"{num}   {title_text}")
        r_num.font.size = Pt(10 if not is_section else 11)
        r_num.bold = is_section
        if is_section:
            r_num.font.color.rgb = RGBColor.from_string("1A237E")

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # SECTION 1: INTRODUCTION
    # ══════════════════════════════════════════════════════════════════════════
    _heading(doc, "SECTION 1    INTRODUCTION", color="1A237E", size=16)
    _add_horizontal_rule(doc, "2E7D32")

    _heading(doc, "1.1 BACKGROUND:", color="1A237E", size=12)
    _body(doc,
        f"{pi.get('client_name','The client')} has engaged UrbanRoof Private Limited to carry out a "
        f"comprehensive Health Assessment / Detailed Diagnostic Report (DDR) for the property located at "
        f"{pi.get('address','the above address')}. The investigation was conducted on {pi.get('inspection_date','the inspection date')} "
        f"by the UrbanRoof technical team using advanced diagnostic tools including IR Thermography, "
        f"Moisture Meter, pH Meter, and Tapping Hammer.")

    _heading(doc, "1.2 OBJECTIVE OF THE HEALTH ASSESSMENT", color="1A237E", size=12)
    objectives = [
        "To detect and document all visible defects, dampness, seepage, cracks, and structural distress.",
        "To correlate visual observations with thermal imaging data for accurate diagnosis of hidden moisture.",
        "To identify probable root causes and leakage travel paths.",
        "To classify defects by severity and recommend appropriate treatment with priority.",
        "To provide a scope-of-work basis for repair estimation and waterproofing treatment design.",
        "To serve as a baseline record for future audits and warranty tracking.",
    ]
    for obj in objectives:
        _bullet(doc, obj)

    _heading(doc, "1.3 SCOPE OF WORK:", color="1A237E", size=12)
    _body(doc,
        f"Visual site inspection using: {pi.get('tools_used', 'Tapping Hammer, Crack Gauge, IR Thermography Camera, Moisture Meter, pH Meter')}. "
        f"Inspection covered all accessible areas including bathrooms, balconies, terraces, external walls, "
        f"RCC members, and internal wall/ceiling surfaces.")

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # SECTION 2: GENERAL INFORMATION
    # ══════════════════════════════════════════════════════════════════════════
    _heading(doc, "SECTION 2    GENERAL INFORMATION", color="1A237E", size=16)
    _add_horizontal_rule(doc, "2E7D32")

    _heading(doc, "2.1 CLIENT & INSPECTION DETAILS", color="1A237E", size=12)
    detail_table = doc.add_table(rows=len(cover_fields), cols=2)
    detail_table.style = "Table Grid"
    for i, (k, v) in enumerate(cover_fields):
        r = detail_table.rows[i]
        _set_cell_bg(r.cells[0], "E8F5E9")
        r.cells[0].text = k
        r.cells[0].paragraphs[0].runs[0].bold = True
        r.cells[0].paragraphs[0].runs[0].font.size = Pt(10)
        r.cells[1].text = str(v)
        r.cells[1].paragraphs[0].runs[0].font.size = Pt(10)
    doc.add_paragraph()

    _heading(doc, "2.2 DESCRIPTION OF SITE", color="1A237E", size=12)
    _body(doc,
        f"The subject property is a {pi.get('property_type','residential property')} comprising {pi.get('floors','multiple')} floor(s), "
        f"approximately {pi.get('age_years','several')} years old. "
        f"Previous audit: {pi.get('previous_audit','Not Available')}. "
        f"Previous repairs: {pi.get('previous_repairs','Not Available')}.")

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # SECTION 3: VISUAL OBSERVATION AND READINGS
    # ══════════════════════════════════════════════════════════════════════════
    _heading(doc, "SECTION 3    VISUAL OBSERVATION AND READINGS", color="1A237E", size=16)
    _add_horizontal_rule(doc, "2E7D32")

    # 3.1 Executive Summary
    _heading(doc, "3.1 EXECUTIVE SUMMARY", color="2E7D32", size=13)
    exec_summary = ddr_json.get("executive_summary") or ddr_json.get("property_issue_summary", "Not Available")
    _body(doc, exec_summary)

    # 3.2 Root Cause / Leakage Source
    _heading(doc, "3.2 SOURCES OF LEAKAGE & ROOT CAUSE", color="2E7D32", size=13)
    _heading(doc, "3.2.1    SUMMARY", color="1A237E", size=11)
    _body(doc, ddr_json.get("property_issue_summary", "Not Available"))

    leakage = ddr_json.get("leakage_sources", "")
    if leakage:
        _heading(doc, "Leakage Travel Path:", color="1A237E", size=11)
        _body(doc, leakage)

    probable = ddr_json.get("probable_root_cause", "")
    if probable:
        _heading(doc, "Probable Root Cause:", color="1A237E", size=11)
        _body(doc, probable)

    doc.add_paragraph()

    # 3.3+ Area-wise observations
    for idx, obs in enumerate(obs_list, start=3):
        area      = obs.get("area", f"Area {idx}")
        floor_lvl = obs.get("floor_level", "")
        heading_prefix = f"3.{idx}"

        section_title = f"{heading_prefix} {area.upper()}"
        if floor_lvl and floor_lvl.lower() != "not available":
            section_title += f"  ({floor_lvl})"
        _heading(doc, section_title, color="2E7D32", size=13)

        # Negative observations
        neg = obs.get("negative_observations", "Not Available")
        _heading(doc, f"{heading_prefix}.1  Negative Side Inputs — {area}", color="1A237E", size=11)
        _body(doc, neg)

        # Structural condition
        struct = obs.get("structural_condition", "")
        if struct and struct.lower() not in ("", "not available"):
            _label_value(doc, "Structural Condition", struct)

        # Moisture readings
        moisture = obs.get("moisture_readings", "")
        if moisture and moisture.lower() not in ("", "not available"):
            _label_value(doc, "Moisture / pH Reading", moisture)

        # Affected surface
        surface = obs.get("affected_surface", "")
        if surface and surface.lower() not in ("", "not available"):
            _label_value(doc, "Affected Surface", surface)

        # ── Inspection images for this area ──
        matched_inspection = _find_images_for_area(inspection_images, area, max_count=2)
        if matched_inspection:
            img_heading = doc.add_paragraph()
            r = img_heading.add_run(f"Inspection Images — {area}")
            r.bold = True; r.font.size = Pt(10)
            r.font.color.rgb = RGBColor.from_string("555555")
            for img in matched_inspection:
                _insert_image(doc, img["path"],
                              caption=f"Fig: {img.get('nearest_heading','')[:80]}",
                              width_inches=4.0)

        # Positive / source side
        pos = obs.get("positive_observations", "")
        if pos and pos.lower() not in ("", "not available"):
            _heading(doc, f"{heading_prefix}.2  Positive / Source Side — {area}", color="1A237E", size=11)
            _body(doc, pos)

        # Thermal readings
        thermal = obs.get("thermal_reading", "")
        if thermal and thermal.lower() not in ("", "not available"):
            _heading(doc, f"{heading_prefix}.3  Thermal Imaging Reading — {area}", color="1A237E", size=11)
            _body(doc, thermal)

            # ── Thermal images for this area ──
            matched_thermal = _find_images_for_area(thermal_images, area, max_count=1, min_bytes=8000)
            if matched_thermal:
                for img in matched_thermal:
                    _insert_image(doc, img["path"],
                                  caption=f"Thermal: {img.get('nearest_heading','')[:80]}",
                                  width_inches=3.5)

        _add_horizontal_rule(doc, "CCCCCC")
        doc.add_paragraph()

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # SECTION 4: SEVERITY ASSESSMENT
    # ══════════════════════════════════════════════════════════════════════════
    _heading(doc, "SECTION 4    SEVERITY ASSESSMENT", color="1A237E", size=16)
    _add_horizontal_rule(doc, "2E7D32")

    severity_list = ddr_json.get("severity_assessment", [])
    if severity_list:
        sev_table = doc.add_table(rows=1, cols=4)
        sev_table.style = "Table Grid"
        for cell, label in zip(sev_table.rows[0].cells, ["Area", "Severity", "Urgency", "Engineering Reasoning"]):
            _set_cell_bg(cell, "1A237E")
            r = cell.paragraphs[0].add_run(label)
            r.bold = True; r.font.size = Pt(10)
            r.font.color.rgb = RGBColor(255, 255, 255)

        sev_colors = {"critical": "FFCDD2", "high": "FFE0B2", "moderate": "FFF9C4", "low": "C8E6C9"}
        for s in severity_list:
            row = sev_table.add_row().cells
            row[0].text = s.get("area", "")
            row[1].text = s.get("severity", "")
            row[2].text = s.get("urgency", "")
            row[3].text = s.get("reasoning", "")
            bg = sev_colors.get(s.get("severity", "").lower(), "FFFFFF")
            for cell in row:
                _set_cell_bg(cell, bg)
                cell.paragraphs[0].runs[0].font.size = Pt(10)
    else:
        _body(doc, "Not Available")

    doc.add_paragraph()
    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # SECTION 5: RECOMMENDED ACTIONS
    # ══════════════════════════════════════════════════════════════════════════
    _heading(doc, "SECTION 5    RECOMMENDED ACTIONS", color="1A237E", size=16)
    _add_horizontal_rule(doc, "2E7D32")

    actions = ddr_json.get("recommended_actions", [])
    if actions:
        act_table = doc.add_table(rows=1, cols=5)
        act_table.style = "Table Grid"
        for cell, label in zip(act_table.rows[0].cells,
                               ["Action", "Area", "Treatment Method", "Priority", "Scope"]):
            _set_cell_bg(cell, "2E7D32")
            r = cell.paragraphs[0].add_run(label)
            r.bold = True; r.font.size = Pt(10)
            r.font.color.rgb = RGBColor(255, 255, 255)

        pri_colors = {"immediate": "FFCDD2", "high": "FFE0B2", "medium": "FFF9C4", "low": "C8E6C9"}
        for action in actions:
            row = act_table.add_row().cells
            row[0].text = action.get("action", "")
            row[1].text = action.get("area", "")
            row[2].text = action.get("treatment_method", "")
            row[3].text = action.get("priority", "")
            row[4].text = action.get("estimated_scope", "")
            bg = pri_colors.get(action.get("priority", "").lower(), "FFFFFF")
            for cell in row:
                _set_cell_bg(cell, bg)
                cell.paragraphs[0].runs[0].font.size = Pt(9)
    else:
        _body(doc, "Not Available")

    doc.add_paragraph()
    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # SECTION 6: WATERPROOFING & STRUCTURAL RECOMMENDATIONS
    # ══════════════════════════════════════════════════════════════════════════
    _heading(doc, "SECTION 6    WATERPROOFING & STRUCTURAL RECOMMENDATIONS", color="1A237E", size=16)
    _add_horizontal_rule(doc, "2E7D32")

    wp = ddr_json.get("waterproofing_recommendations", "")
    if wp:
        _heading(doc, "6.1 Waterproofing System Recommendations", color="2E7D32", size=12)
        _body(doc, wp)

    struct = ddr_json.get("structural_recommendations", "")
    if struct:
        _heading(doc, "6.2 Structural Repair Recommendations", color="2E7D32", size=12)
        _body(doc, struct)

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # SECTION 7: ADDITIONAL NOTES
    # ══════════════════════════════════════════════════════════════════════════
    _heading(doc, "SECTION 7    ADDITIONAL NOTES", color="1A237E", size=16)
    _add_horizontal_rule(doc, "2E7D32")
    _body(doc, ddr_json.get("additional_notes", "None"))

    missing = ddr_json.get("missing_or_unclear_info", "")
    if missing:
        _heading(doc, "Missing / Unclear Information:", color="C62828", size=12)
        _body(doc, missing)

    doc.add_paragraph()
    footer_p = doc.add_paragraph()
    footer_p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    fr = footer_p.add_run(
        "UrbanRoof Private Limited  ·  Detailed Diagnostic Report  ·  Strictly Confidential")
    fr.font.size = Pt(8); fr.italic = True
    fr.font.color.rgb = RGBColor.from_string("888888")

    doc.save(output_path)
    return output_path